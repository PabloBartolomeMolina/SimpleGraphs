"""
File to manage the creation of a report file in Word format.
Purpose is to create a file, with some fixed sections and information and using the graphs and others with the
data extracted from Excel file.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Creation of the document.
def create_document():
    document = Document()
    document.add_heading('Demo heading, level 1', level=1)
    document.save('demo.docx')


# Open a given document if it exists.
def open_document(filename):
    try:
        document = Document(f'./{filename}.docx')
        print("Document: "f'{filename}.docx' " is open")
        document.add_paragraph(
            'First item in ordered list', style='List Number'
        )
        document.add_paragraph(
            'Second item in ordered list', style='List Number'
        )
        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )
        table = document.add_table(rows=2, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = 250000
        for qty, id1, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id1
            row_cells[2].text = desc

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p = document.add_paragraph()
        r = p.add_run()
        r.add_text('Good Morning every body,This is my graph : ')
        r.add_picture('./simplegraph1.png')
        r.add_text("Demo image with demo data")

        document.save('demo.docx')

        # Another document. Add a table with an image inside it.
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.style = 'insideH'
        cell = table.cell(0, 1)
        cell.text = 'parrot,\npossibly dead'
        row = table.rows[1]
        row.cells[0].text = 'Foo bar to you.'
        row.cells[1].text = 'And a hearty foo bar to you too sir!'
        paragraph = table.cell(0, 0).paragraphs[0]
        run = paragraph.add_run()
        run.add_picture('./simplegraph1.png', width=2500000, height=2000000)

        doc.save('addImage.docx')
    except:
        print("\nDocument not found:\n" f'{filename}.docx')


# Function to generate a proper report
def true_report():
    pass
